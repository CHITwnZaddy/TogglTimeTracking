"""
Generate the Toggl Invoice Generator technical spec as a Word document.

Usage:
    python3 docs/generate_spec.py

Output:
    ~/Downloads/Toggl_Invoice_Generator_Spec.docx

Edit the content below and re-run to regenerate.
"""
from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH


OUTPUT = Path.home() / "Downloads" / "Toggl_Invoice_Generator_Spec.docx"


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x1F, 0x2A, 0x44)  # dark navy


def add_paragraph(doc: Document, text: str, bold: bool = False) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(11)
    if bold:
        run.bold = True


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(item)
        run.font.size = Pt(11)


def add_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Light Grid Accent 1"

    # Header row
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for para in cell.paragraphs:
            for run in para.runs:
                run.bold = True
                run.font.size = Pt(11)

    # Data rows
    for r, row in enumerate(rows, start=1):
        for c, val in enumerate(row):
            cell = table.rows[r].cells[c]
            cell.text = val
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.size = Pt(10)

    doc.add_paragraph()  # spacer


def add_code(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Menlo"
    run.font.size = Pt(10)


def build() -> None:
    doc = Document()

    # ── Default style ────────────────────────────────────────────────────────
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    # ── Title ────────────────────────────────────────────────────────────────
    title = doc.add_heading("Toggl Invoice Generator — Technical Specification", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT

    meta = doc.add_paragraph()
    meta.add_run("Author: ").bold = True
    meta.add_run("Austin Alexander Guzman\n")
    meta.add_run("Last updated: ").bold = True
    meta.add_run("April 2026\n")
    meta.add_run("Live app: ").bold = True
    meta.add_run("https://toggltimetracking-aag.streamlit.app/\n")
    meta.add_run("Repository: ").bold = True
    meta.add_run("https://github.com/AustinAlexanderGuzman/TogglTimeTracking")

    # ── 1. Overview ──────────────────────────────────────────────────────────
    add_heading(doc, "1. Overview", level=1)
    add_paragraph(
        doc,
        "The Toggl Invoice Generator is a private, single-user web app that turns "
        "time-tracking data from Toggl Track into formatted Excel invoices and emails "
        "them directly to clients. It replaces a manual workflow — exporting CSVs from "
        "Toggl, copying them into a template spreadsheet, and composing client emails "
        "by hand — with a three-click process: fetch, generate, send."
    )
    add_paragraph(
        doc,
        "The app is deployed on Streamlit Community Cloud behind a password plus "
        "time-based one-time password (TOTP) login gate. Invoices are generated as "
        ".xlsx files that exactly match the consulting firm's legacy template, and "
        "emails are delivered via Google Workspace SMTP so they arrive from the "
        "user's real work address rather than a third-party service."
    )

    # ── 2. Architecture ──────────────────────────────────────────────────────
    add_heading(doc, "2. Architecture", level=1)
    add_paragraph(
        doc,
        "The app is organized into one UI module and three service modules, each "
        "handling a single concern. This keeps app.py focused on presentation and "
        "state while isolating external integrations and file generation."
    )

    add_table(
        doc,
        headers=["Module", "Responsibility", "Depends on"],
        rows=[
            ["app.py", "Streamlit UI, login gate, session state, orchestration", "All three service modules"],
            ["toggl_client.py", "Toggl Track API v9 calls, response enrichment", "requests"],
            ["invoice_generator.py", "Excel (.xlsx) generation with exact template formatting", "openpyxl"],
            ["email_sender.py", "SMTP delivery of .xlsx attachment via Google Workspace", "smtplib (stdlib)"],
        ],
    )

    add_paragraph(doc, "Data flow (happy path):", bold=True)
    add_bullets(
        doc,
        [
            "User logs in (password + TOTP) → session flagged authenticated.",
            "User selects date range → app.py calls toggl_client.get_enriched_entries().",
            "toggl_client fetches time entries, projects, and tasks, joins them, returns enriched records.",
            "User filters by project and clicks Generate → app.py calls invoice_generator.generate_invoice().",
            "invoice_generator returns raw .xlsx bytes → app.py stores them in st.session_state.",
            "User clicks Download (serves bytes) or Send Email → app.py calls email_sender.send_invoice_email().",
        ],
    )

    # ── 3. Tech Stack ────────────────────────────────────────────────────────
    add_heading(doc, "3. Tech Stack", level=1)
    add_table(
        doc,
        headers=["Component", "Choice", "Rationale"],
        rows=[
            ["Language", "Python 3.9+", "Mature ecosystem for data/Excel work; matches host runtime on Streamlit Cloud."],
            ["UI framework", "Streamlit >= 1.35", "Zero-config web UI; perfect for internal single-user tools."],
            ["HTTP client", "requests", "De facto Python HTTP library; simple Basic Auth for Toggl API."],
            ["Excel library", "openpyxl", "Full control over fonts, fills, merged cells, number formats."],
            ["MFA", "pyotp", "RFC 6238 TOTP; compatible with Google Authenticator, 1Password, Authy."],
            ["Email", "smtplib + email.message (stdlib)", "No new runtime dependency; Gmail SMTP needs nothing more."],
            ["Secrets (local)", "python-dotenv + .streamlit/secrets.toml", "Local dev reads either; committed .env is an anti-pattern."],
            ["Secrets (cloud)", "Streamlit Cloud Secrets UI", "st.secrets injects values at runtime; never in the repo."],
            ["Hosting", "Streamlit Community Cloud", "Free for public apps; auto-deploys on git push; persistent WebSocket."],
        ],
    )

    # ── 4. Authentication ────────────────────────────────────────────────────
    add_heading(doc, "4. Authentication", level=1)
    add_paragraph(
        doc,
        "Because the app is publicly reachable at a streamlit.app URL, access is "
        "gated by two factors on every session: a static password and a time-based "
        "one-time password (TOTP)."
    )

    add_paragraph(doc, "Login flow:", bold=True)
    add_bullets(
        doc,
        [
            "login_screen() renders a form asking for password + 6-digit code.",
            "Password is compared to APP_PASSWORD from secrets.",
            "TOTP code is verified with pyotp.TOTP(TOTP_SECRET).verify(code, valid_window=1).",
            "valid_window=1 accepts the current and adjacent 30-second windows, forgiving minor clock drift.",
            "On success, st.session_state['authenticated'] = True and st.rerun() loads the main UI.",
            "Until authenticated is True, st.stop() prevents any downstream code from running.",
        ],
    )

    add_paragraph(
        doc,
        "The TOTP secret is a base32 string generated once and added to the "
        "authenticator app (e.g., Google Authenticator). The app never stores "
        "generated codes — it only verifies that the submitted code matches the "
        "current time window derived from TOTP_SECRET."
    )

    # ── 5. Toggl Integration ─────────────────────────────────────────────────
    add_heading(doc, "5. Toggl Integration", level=1)
    add_paragraph(
        doc,
        "The app uses Toggl Track API v9 with HTTP Basic Auth. The API token "
        "(found in Toggl profile settings) is sent as the username, with the "
        "literal string 'api_token' as the password."
    )

    add_table(
        doc,
        headers=["Endpoint", "Purpose", "Called from"],
        rows=[
            ["GET /api/v9/me", "Fetch current user and workspace IDs", "TogglClient.get_user()"],
            ["GET /api/v9/me/time_entries", "Fetch time entries in a date range", "TogglClient.get_time_entries()"],
            ["GET /api/v9/workspaces/{wid}/projects/{pid}", "Resolve project name from project_id", "TogglClient._get_project()"],
            ["GET /api/v9/workspaces/{wid}/tasks/{tid}", "Resolve task name from task_id", "TogglClient._get_task()"],
        ],
    )

    add_paragraph(doc, "Enrichment logic:", bold=True)
    add_bullets(
        doc,
        [
            "Raw time entries only contain project_id and task_id integers, not names.",
            "get_enriched_entries() calls _get_project() and _get_task() per entry to resolve names.",
            "Running timers (duration < 0) are skipped — only completed entries are billable.",
            "Results are sorted by start time ascending for consistent invoice ordering.",
        ],
    )

    # ── 6. Invoice Generation ────────────────────────────────────────────────
    add_heading(doc, "6. Invoice Generation", level=1)
    add_paragraph(
        doc,
        "generate_invoice() produces a byte-for-byte compatible replica of the "
        "consulting firm's legacy .xlsx invoice template, then returns the raw "
        "bytes so Streamlit can serve them as a download or attach them to an email."
    )

    add_paragraph(doc, "Template fidelity rules:", bold=True)
    add_bullets(
        doc,
        [
            "Font: Actor Regular 12pt (bold and regular variants).",
            "Row 1 (INVOICE #), Row 3 (PERIOD): label in column A, value merged across B1:I1 / B3:I3.",
            "Row 5: column headers (Start date, Stop date, Project, Task, Description, Duration, Member, Amount).",
            "Rows 6+: data rows, one per time entry.",
            "Last row: total row with 'Total' labels, aggregate duration, and total amount.",
            "Row heights: 26 for info rows, 27 for data/total rows.",
            "Column widths match the original template to the hundredth (e.g., A=10.83, E=32.83).",
            "Date cells use number format 'DD MMM YY' (e.g., 22 APR 26).",
            "Duration cells use 'h:mm;@' for individual rows.",
            "Amount cells use '#,##0.00\\ \"USD\"' for currency formatting.",
        ],
    )

    add_paragraph(doc, "Visual formatting additions:", bold=True)
    add_bullets(
        doc,
        [
            "Header row (row 5): bold, centered, background #C3C4EB.",
            "Data rows: alternating backgrounds — #FFFFFF (even) and #EFF0FA (odd).",
            "Start date column (B): bold on every data row.",
            "Amount column (I): bold on every data row.",
            "Total row: background #E1E1F5; all 'Total' labels and total amount in bold.",
        ],
    )

    add_paragraph(doc, "Quarter-hour rounding (total Duration only):", bold=True)
    add_paragraph(
        doc,
        "The original approach used an Excel SUM formula on datetime.time values, but "
        "Excel stores time as fractions of a day, which caused display bugs "
        "(e.g., 5h 20m rendering as ':20'). The fix: calculate total hours in Python, "
        "then round to the nearest 0.25 using the formula "
        "round(total_seconds / 3600 * 4) / 4. The result is written as a plain "
        "decimal number with format 0.00\" hrs\". Individual row durations still use "
        "h:mm since those are exact, not billed."
    )

    add_paragraph(doc, "Invoice numbering:", bold=True)
    add_bullets(
        doc,
        [
            "Format: MMM YY - PROJECT - NNNN (e.g., APR 26 - STL - 1001).",
            "Counter stored per-project in invoice_counter.json; increments by 1 each generation.",
            "Auto-fills only when exactly one project is selected (multi-project invoices need manual numbering).",
            "invoice_counter.json is gitignored; on Streamlit Cloud's read-only filesystem, save_counters() catches OSError and continues (counter resets gracefully on redeploy).",
        ],
    )

    # ── 7. Email Delivery ────────────────────────────────────────────────────
    add_heading(doc, "7. Email Delivery", level=1)
    add_paragraph(
        doc,
        "Invoices are emailed directly from the app via Google Workspace SMTP "
        "(smtp.gmail.com:587, STARTTLS). The From address is the user's real "
        "work email, so recipients don't see third-party service tags like "
        "'via sendgrid.net' that can trigger spam filtering or erode trust."
    )

    add_paragraph(doc, "Why SMTP + App Password over OAuth or SendGrid:", bold=True)
    add_bullets(
        doc,
        [
            "Simplicity: no OAuth dance, no token refresh, no webhook for delivery status.",
            "Zero new runtime deps: smtplib and email.message are in Python's stdlib.",
            "Deliverability: messages land in the user's Gmail Sent folder automatically.",
            "Cost: free — no third-party email service bill.",
        ],
    )

    add_paragraph(doc, "Setup (one-time):", bold=True)
    add_bullets(
        doc,
        [
            "Enable 2-Step Verification on the Google account.",
            "Generate an App Password at myaccount.google.com/apppasswords (16-char code).",
            "Store as SMTP_USER (the Google account email) and SMTP_PASSWORD (the App Password) in secrets.",
            "Note: App Passwords are account-scoped. A CMP-workspace App Password cannot authenticate a personal @gmail.com account, even with the same browser.",
        ],
    )

    add_paragraph(doc, "Per-project client mapping:", bold=True)
    add_paragraph(
        doc,
        "client_emails.json is a flat JSON map of project-code → client email "
        "(e.g., {\"STL\": \"carrine@cmpstl.com\"}). When exactly one project is "
        "selected, the To field auto-fills from this map. The Cc field defaults to "
        "SMTP_USER (copy to self). The file is gitignored — it contains PII and "
        "should never enter the repo."
    )

    add_paragraph(doc, "Session state for reliability:", bold=True)
    add_paragraph(
        doc,
        "Streamlit reruns the full script on every widget interaction. Without "
        "session state, the .xlsx bytes returned by generate_invoice() would be "
        "regenerated (and re-number) on every email-form keystroke. The fix: the "
        "Generate button stores xlsx, filename, invoice_number, total_amount, and "
        "period dates in st.session_state. Download and Send Email both read from "
        "session state, guaranteeing they operate on the same generated artifact."
    )

    # ── 8. Deployment ────────────────────────────────────────────────────────
    add_heading(doc, "8. Deployment", level=1)

    add_paragraph(doc, "Local development:", bold=True)
    add_bullets(
        doc,
        [
            "Python 3.9+ with packages from requirements.txt installed via pip3 install --user.",
            "Streamlit binary lives at ~/Library/Python/3.9/bin/streamlit — not on default zsh PATH.",
            "A Desktop .command launcher exports the correct PATH and runs streamlit run app.py in a Terminal window for one-click launch.",
            "Secrets live in .streamlit/secrets.toml (gitignored) and are read via st.secrets.",
        ],
    )

    add_paragraph(doc, "Cloud deployment (Streamlit Community Cloud):", bold=True)
    add_bullets(
        doc,
        [
            "Git push to main auto-triggers a redeploy.",
            "Main file path must be set to app.py in the app's Settings (not toggl_client.py — this caused an early blank-screen bug).",
            "Secrets configured via the Streamlit Cloud UI, not committed to the repo.",
            "requirements.txt drives package installation on the cloud runtime.",
            ".devcontainer/devcontainer.json mirrors the Streamlit command for consistent environments.",
        ],
    )

    # ── 9. Configuration Reference ───────────────────────────────────────────
    add_heading(doc, "9. Configuration Reference", level=1)
    add_paragraph(doc, "All secrets are set in both .streamlit/secrets.toml (local) and the Streamlit Cloud Secrets UI (production).")

    add_table(
        doc,
        headers=["Key", "Type", "Description"],
        rows=[
            ["TOGGL_API_TOKEN", "str", "Toggl Track API token from Profile Settings > Profile > API Token."],
            ["HOURLY_RATE", "str (parsed as float)", "Default billing rate in USD; editable per session in the sidebar."],
            ["APP_PASSWORD", "str", "Static password for the login gate."],
            ["TOTP_SECRET", "str (base32)", "Shared secret for TOTP MFA; add to authenticator app once."],
            ["SMTP_USER", "str", "Google Workspace email used as both SMTP login and From address."],
            ["SMTP_PASSWORD", "str (16 chars)", "Google App Password generated for the SMTP_USER account."],
        ],
    )

    add_paragraph(doc, "Non-secret config files:", bold=True)
    add_table(
        doc,
        headers=["File", "Committed?", "Purpose"],
        rows=[
            ["requirements.txt", "Yes", "Runtime Python dependencies."],
            ["invoice_counter.json", "No (gitignored)", "Per-project auto-increment counter state."],
            ["client_emails.json", "No (gitignored, PII)", "Project-code to client-email map for auto-filling To field."],
            [".gitignore", "Yes", "Excludes secrets, counter, client map, .DS_Store, node_modules."],
            [".devcontainer/devcontainer.json", "Yes", "Cloud runtime entry point."],
        ],
    )

    # ── 10. Known Limitations / Future Ideas ─────────────────────────────────
    add_heading(doc, "10. Known Limitations and Future Ideas", level=1)

    add_paragraph(doc, "Current limitations:", bold=True)
    add_bullets(
        doc,
        [
            "Single-user: no multi-user auth model; one APP_PASSWORD and one TOTP_SECRET guard the whole app.",
            "Invoice counter resets on Streamlit Cloud redeploy (filesystem is ephemeral). Local runs persist; cloud runs do not.",
            "client_emails.json is hand-edited — no UI to add or update client mappings.",
            "Toggl API calls are sequential, not batched. Large date ranges with many distinct projects/tasks produce many HTTP requests.",
            "Multi-project invoices require manual invoice numbering (auto-number only fires when exactly one project is selected).",
            "All times are displayed in UTC as stored in Toggl. No local timezone conversion.",
        ],
    )

    add_paragraph(doc, "Deferred ideas (in priority order):", bold=True)
    add_bullets(
        doc,
        [
            "Sidebar 'Manage client emails' expander to add, edit, and remove mappings from the UI.",
            "Cache project/task lookups per session to reduce Toggl API calls.",
            "Persist invoice_counter.json to an external store (S3, Gist, or a free tier KV store) so cloud counters survive redeploys.",
            "Per-client hourly rates stored alongside the email map.",
            "Timezone setting in sidebar with automatic conversion of entry start/stop times.",
            "PDF export alongside .xlsx for clients who prefer non-editable invoices.",
            "Read-receipt or delivery-confirmation webhook so the user knows the client opened the email.",
        ],
    )

    # ── Footer ───────────────────────────────────────────────────────────────
    doc.add_paragraph()
    footer = doc.add_paragraph()
    footer_run = footer.add_run(
        "This spec is generated from docs/generate_spec.py. "
        "Edit the Python source and re-run to regenerate."
    )
    footer_run.italic = True
    footer_run.font.size = Pt(9)
    footer_run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    # ── Save ─────────────────────────────────────────────────────────────────
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(f"Wrote {OUTPUT}")


if __name__ == "__main__":
    build()
